"""
Gera o .docx da Mesa Redonda Agentica - 17/04/2026
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

AZUL = RGBColor(0x00, 0x10, 0x70)
CINZA = RGBColor(0x44, 0x44, 0x44)
PRETO = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# --- Margens 2.5cm ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Estilos base ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = PRETO

def add_title(text, size=16, bold=True, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = AZUL
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = AZUL
    run.font.name = 'Calibri'
    return p

def add_subsection(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = AZUL
    run.font.name = 'Calibri'
    return p

def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    if level == 1:
        p.paragraph_format.left_indent = Cm(1.5)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.name = 'Calibri'
        run_n = p.add_run(text)
        run_n.font.size = Pt(11)
        run_n.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    return p

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

# ===================================================================
# CAPA
# ===================================================================
doc.add_paragraph()  # espacamento
add_title("THAUMA", size=24, space_after=2)
add_title("Inteligencia & Narrativa em Saude", size=12, bold=False, space_after=20)

add_title("DOCUMENTO ESTRATEGICO", size=16, space_after=4)
add_title("MESA REDONDA AGENTICA", size=18, space_after=4)
add_title("ABRIL 2026", size=14, bold=False, space_after=20)

add_title("Comunidade como Canal de Conversao:", size=13, bold=True, space_after=2)
add_title("Estrategia ou Dispersao?", size=13, bold=False, space_after=30)

# Metadados
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(2)
for line in [
    "Provocacao: Pedro William Ribeiro Diniz (Fundador)",
    "Conducao: Socrates (CEO)",
    "Data: 17/04/2026",
    "",
    "Participantes:",
    "Pericles (Marketing) | Arquimedes (Projetos) | Tales (Financeiro)",
    "Icaro (Novos Produtos) | Hefesto (Operacoes)",
]:
    if line == "":
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_after = Pt(2)
        continue
    if meta.text:
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_after = Pt(2)
    run = meta.add_run(line)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = CINZA

# Classificacao
doc.add_paragraph()
cls_p = doc.add_paragraph()
cls_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cls_p.add_run("CLASSIFICACAO: ESTRATEGICO-DECISORIO | GOVERNANCA INTERNA")
run.bold = True
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.font.color.rgb = AZUL

doc.add_page_break()

# ===================================================================
# 1. SUMARIO EXECUTIVO
# ===================================================================
add_section_heading("1. SUMARIO EXECUTIVO")

add_body("Decisoes-chave desta Mesa Redonda:", bold=True)

sumario_items = [
    ("S1. ", "Comunidade NAO e prioridade agora - nao iniciar antes do gate F0 do Higia (15/05)."),
    ("S2. ", "O problema de conversao e cadencia de follow-up, nao canal."),
    ("S3. ", "Trilha A (S1-S4) como teste de 4 semanas antes de qualquer canal novo."),
    ("S4. ", "WhatsApp descartado como plataforma (unanimidade 5/5)."),
    ("S5. ", "Newsletter Aletheia como ponte logica entre conteudo e comunidade."),
    ("S6. ", "Revisitar no Conselho #2 pos-F0 (19/05)."),
    ("S7. ", "Pedro nao sera operador de comunidade em hipotese alguma."),
]
for prefix, text in sumario_items:
    add_bullet(text, bold_prefix=prefix)

add_separator()

# ===================================================================
# 2. CONTEXTO E PROVOCACAO
# ===================================================================
add_section_heading("2. CONTEXTO E PROVOCACAO")

add_body(
    "Pedro questionou se deveria construir uma comunidade (WhatsApp ou similar) "
    "a partir dos seguidores do LinkedIn e dos leads gerados por lead magnets, "
    "com o objetivo de aumentar conversao. O proprio Pedro levantou a hipotese "
    "de que isso poderia ser escapismo - uma forma de evitar o trabalho duro de "
    "follow-up direto."
)

add_subsection("Situacao atual:")
situacao = [
    "1 cliente ativo: Santa Casa de Ouro Preto (SCOP) - R$ 18.997 (1a parcela recebida)",
    "16 leads no pipeline: 3 quentes, 6 mornos, 7 frios",
    "Lead magnets geraram 5+ leads com ZERO conversoes ate o momento",
    "Disponibilidade de Pedro: 15 horas/semana para THAUMA",
    "Higia F0 (gate de validacao) em 28 dias (15/05/2026)",
]
for item in situacao:
    add_bullet(item)

add_body(
    "A pergunta central: dada a restricao de tempo e a fase atual da empresa, "
    "construir comunidade e alavanca ou dispersao?",
    italic=True
)

add_separator()

# ===================================================================
# 3. PARTICIPANTES
# ===================================================================
add_section_heading("3. PARTICIPANTES DA MESA REDONDA")

# Table
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Shading Accent 1'

headers = ["Participante", "Departamento", "Papel", "Presenca"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = 'Calibri'

data = [
    ("Socrates", "CEO", "Conducao", "OBRIGATORIO"),
    ("Pericles", "Marketing", "Posicao", "OBRIGATORIO"),
    ("Arquimedes", "Projetos", "Posicao", "OBRIGATORIO"),
    ("Tales", "Financeiro", "Posicao", "OBRIGATORIO"),
    ("Icaro", "Novos Produtos", "Posicao", "CONVIDADO"),
    # Hefesto na row 5 (index)
]
# Adjust: 6 rows total (1 header + 5 data). Let me fix.
# Actually we have 5 participants + header = 6 rows. But Socrates is row 1.
# Wait, we have: Socrates, Pericles, Arquimedes, Tales, Icaro, Hefesto = 6 participants + 1 header = 7 rows

# Recreate table
# Remove old table and redo
doc.element.body.remove(table._element)

table = doc.add_table(rows=7, cols=4)
table.style = 'Light Shading Accent 1'

headers = ["Participante", "Departamento", "Papel", "Presenca"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = 'Calibri'

data = [
    ("Socrates", "CEO", "Conducao", "PERMANENTE"),
    ("Pericles", "Marketing", "Posicao", "OBRIGATORIO"),
    ("Arquimedes", "Projetos", "Posicao", "OBRIGATORIO"),
    ("Tales", "Financeiro", "Posicao", "OBRIGATORIO"),
    ("Icaro", "Novos Produtos", "Posicao", "CONVIDADO"),
    ("Hefesto", "Operacoes", "Posicao", "CONVIDADO"),
]
for row_idx, (nome, dept, papel, presenca) in enumerate(data, start=1):
    table.rows[row_idx].cells[0].text = nome
    table.rows[row_idx].cells[1].text = dept
    table.rows[row_idx].cells[2].text = papel
    table.rows[row_idx].cells[3].text = presenca
    for cell in table.rows[row_idx].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = 'Calibri'

add_separator()

# ===================================================================
# 4. POSICOES DOS GERENTES
# ===================================================================
add_section_heading("4. POSICOES DOS GERENTES")

# Pericles
add_subsection("4.1 Pericles (Marketing) - CONDICIONAL")
add_body(
    "Voto: Comunidade SIM, mas nao agora. Problema real, diagnostico errado.",
    bold=True
)
add_body(
    "Pericles argumentou que o problema de conversao e real, mas o diagnostico esta errado. "
    "O gargalo nao e falta de canal - e falta de cadencia no follow-up dos leads existentes. "
    "Os 16 leads no pipeline representam potencial inexplorado. Lead magnets geraram contatos "
    "que nunca receberam uma segunda interacao estruturada."
)
add_body("Posicao: Cadencia > Canal. Comunidade so apos F0 + 4 semanas de Trilha A executada.", italic=True)

# Arquimedes
add_subsection("4.2 Arquimedes (Projetos) - CONTRA")
add_body(
    "Voto: Nao. Dispersao classica.",
    bold=True
)
add_body(
    "Arquimedes foi enfatico: com 15 horas semanais e entregas pendentes (SCOP, Higia F0), "
    "adicionar um canal novo e dispersao. Invocou o principio 'um arquetipo por vez' - "
    "a THAUMA ainda nao dominou o arquetipo de venda consultiva direta. Comunidade e um "
    "arquetipo diferente que exige infraestrutura, moderacao e conteudo proprio. "
    "Ha pendencias atrasadas que devem ser priorizadas."
)

# Tales
add_subsection("4.3 Tales (Financeiro) - CONTRA")
add_body(
    "Voto: Nao. Matematica nao fecha.",
    bold=True
)
add_body(
    "Tales apresentou a analise de custo de oportunidade:"
)
add_bullet("Hora de Pedro em follow-up direto: potencial de R$ 1.500 a R$ 3.250 por hora")
add_bullet("Hora de Pedro em comunidade: potencial de R$ 0 a R$ 40 por hora")
add_bullet("Cobranca da 2a parcela SCOP (1 hora de trabalho) = R$ 9.498,50 de receita")
add_body(
    "Conclusao de Tales: cada hora gasta em comunidade em vez de follow-up/cobranca "
    "e destruicao de valor. A matematica e inequivoca.",
    italic=True
)

# Icaro
add_subsection("4.4 Icaro (Novos Produtos) - CONDICIONAL")
add_body(
    "Voto: Sim, como extensao da Aletheia, pos-F0, com dono que nao seja Pedro.",
    bold=True
)
add_body(
    "Icaro trouxe a perspectiva de produto: comunidade pode funcionar como extensao "
    "natural da newsletter Aletheia (que ja esta no roadmap). O fluxo seria: "
    "LinkedIn -> Lead Magnet -> Newsletter Aletheia -> Comunidade. "
    "Mas alertou para o risco de posicionamento: comunidade mal executada pode "
    "diminuir a percepcao de exclusividade que o Prisma exige."
)

# Hefesto
add_subsection("4.5 Hefesto (Operacoes) - CONTRA")
add_body(
    "Voto: Nao. Infraestrutura nao existe e custo operacional e proibitivo.",
    bold=True
)
add_body(
    "Hefesto alertou que a infraestrutura necessaria para uma comunidade funcional "
    "simplesmente nao existe hoje. Estimativa de 4 a 7 horas semanais sem automacao "
    "para moderacao, conteudo e engajamento. As automacoes pendentes de integracao "
    "(Notion, Obsidian, n8n) ja estao atrasadas. Adicionar mais um sistema seria "
    "insustentavel com a capacidade atual."
)

add_separator()

# ===================================================================
# 5. CONSENSOS
# ===================================================================
add_section_heading("5. CONSENSOS FORMADOS")

consensos = [
    ("C1", "Comunidade NAO e prioridade agora", "Unanimidade. Nao iniciar antes do gate F0 do Higia (15/05)."),
    ("C2", "Problema e cadencia, nao canal", "Votacao 5/5. O gargalo de conversao esta na falta de follow-up estruturado, nao na ausencia de um canal comunitario."),
    ("C3", "Executar Trilha A antes de canal novo", "Trilha A (follow-up estruturado S1-S4) deve ser executada e medida antes de considerar qualquer canal adicional."),
    ("C4", "WhatsApp descartado", "Votacao 5/5. Plataforma inadequada para o posicionamento premium da THAUMA. Risco de informalidade excessiva."),
    ("C5", "Aletheia como ponte", "Newsletter Aletheia sera o elo entre conteudo LinkedIn e eventual comunidade. 1a edicao prevista para 31/05."),
    ("C6", "Revisitar no Conselho #2", "Tema sera reavaliado no Conselho Estrategico #2, apos gate F0 do Higia (19/05)."),
    ("C7", "Pedro nao sera operador", "Em qualquer cenario futuro, Pedro nao assume operacao direta de comunidade. Papel: curadoria estrategica apenas."),
]

for cod, titulo, desc in consensos:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run_cod = p.add_run(f"{cod} - {titulo}")
    run_cod.bold = True
    run_cod.font.size = Pt(11)
    run_cod.font.name = 'Calibri'
    run_cod.font.color.rgb = AZUL
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.left_indent = Cm(0.5)
    run_desc = p2.add_run(desc)
    run_desc.font.size = Pt(11)
    run_desc.font.name = 'Calibri'

add_separator()

# ===================================================================
# 6. PLANO DE ACAO
# ===================================================================
add_section_heading("6. PLANO DE ACAO")

# Table
action_table = doc.add_table(rows=6, cols=4)
action_table.style = 'Light Shading Accent 1'

headers = ["Quando", "Acao", "Responsavel", "Status"]
for i, h in enumerate(headers):
    cell = action_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = 'Calibri'

actions = [
    ("17/04/2026", "Registrar decisao e atualizar Obsidian", "Socrates", "IMEDIATO"),
    ("17/04 - 15/05", "Trilha A (S1-S4): follow-up estruturado\nTrilha B: conteudo LinkedIn", "Pericles / Pedro", "EM EXECUCAO"),
    ("31/05/2026", "Newsletter Aletheia #1", "Pericles", "PLANEJADO"),
    ("19/05/2026", "Conselho #2: reavaliar comunidade", "Socrates", "AGENDADO"),
    ("Condicional", "Blueprint de comunidade (se decisao GO no Conselho #2)", "Icaro / Pericles", "CONDICIONAL"),
]

for row_idx, (quando, acao, resp, status) in enumerate(actions, start=1):
    action_table.rows[row_idx].cells[0].text = quando
    action_table.rows[row_idx].cells[1].text = acao
    action_table.rows[row_idx].cells[2].text = resp
    action_table.rows[row_idx].cells[3].text = status
    for cell in action_table.rows[row_idx].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = 'Calibri'

add_separator()

# ===================================================================
# 7. NOTA DE SOCRATES
# ===================================================================
add_section_heading("7. NOTA DE SOCRATES")

add_subsection("Confiante:")
add_body(
    "Pedro diagnosticou o padrao de escapismo antes de agir. Isso e maturidade estrategica. "
    "Trazer a pergunta para a Mesa Redonda em vez de simplesmente executar demonstra que o "
    "metodo maieutico esta internalizando. A decisao 'nao agora' nao e 'nao nunca' - "
    "e disciplina de priorizacao."
)

add_subsection("Em guarda:")
add_body(
    "A Trilha A existe ha 5 dias sem execucao. O risco nao e decidir errado sobre comunidade - "
    "e nao executar o que ja foi decidido sobre follow-up. O antidoto para procrastinacao "
    "nao e mais planejamento. E a primeira ligacao.",
    bold=True
)

doc.add_paragraph()
doc.add_paragraph()

# Assinatura
sig = doc.add_paragraph()
sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sig.add_run("Socrates - CEO | THAUMA Inteligencia & Narrativa em Saude")
run.italic = True
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = AZUL

sig2 = doc.add_paragraph()
sig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sig2.add_run("\"O espanto da descoberta. A ciencia do resultado.\"")
run2.italic = True
run2.font.size = Pt(9)
run2.font.name = 'Calibri'
run2.font.color.rgb = CINZA

# ===================================================================
# SALVAR
# ===================================================================
output_dir = r"C:\Users\User\thauma-staff\docs_internos"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "THAUMA_Mesa_Redonda_Comunidade_2026-04-17.docx")
doc.save(output_path)
print(f"Documento salvo em: {output_path}")
print(f"Tamanho: {os.path.getsize(output_path):,} bytes")
